# Main entry point for FASAPI backend API

from fastapi import FastAPI, Header, Body, Response, HTTPException, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional, Any, Dict
from pydantic import BaseModel
from datetime import datetime, timezone, timedelta
from uuid import uuid4
import logging
import sys
import json
import base64
import io
import re
import os
import markdown
from bs4 import BeautifulSoup
import subprocess
import tempfile
import shutil

# Document generation imports
from docx import Document
import pandas as pd

# import modules that might use print() - AFTER redirection is set up
from utils.terminal_log_handler import BlobLogHandler
from config.decryption import decode
from agents.supervisor import supervisor
from agents.contextllm import rewrite_question
from tools.download_detector import detect_download_request  # for early detection of download intent in user query


# In-memory stores with timestamps for expiration
MESSAGES: Dict[str, Dict[str, Any]] = {}
TABLES: Dict[str, Dict[str, Any]] = {}
LAST_IDS_BY_SESSION: Dict[str, Dict[str, Optional[str]]] = {}


# Redirect print() to logging so all terminal output is captured
class PrintLogger:
    def __init__(self, logger, level=logging.INFO):
        self.logger = logger
        self.level = level
        self.terminal = sys.stdout  # Keep original stdout
    
    def write(self, message):
        if message.strip():  # Avoid empty lines
            self.logger.log(self.level, message.strip())
        self.terminal.write(message)  # Also write to console
    
    def flush(self):
        self.terminal.flush()


# Configure logging FIRST before creating any loggers
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),  # Keep console output
    ]
)

# Redirect print() to logging BEFORE importing modules that use print()
sys.stdout = PrintLogger(logging.getLogger('stdout'))


# Add blob storage handler ONLY in production (not LOCAL)
IS_PRODUCTION = os.getenv("ENVIRONMENT", "LOCAL").upper() != "LOCAL"

if IS_PRODUCTION:
    try:
        blob_handler = BlobLogHandler()
        blob_handler.setLevel(logging.INFO)
        blob_handler.setFormatter(logging.Formatter('%(name)s - %(levelname)s - %(message)s'))
        logging.getLogger().addHandler(blob_handler)
        print("Blob log handler initialized successfully (PRODUCTION MODE)")
    except Exception as e:
        print(f"Failed to initialize blob log handler: {e}")
else:
    print("Blob log handler DISABLED (LOCAL DEVELOPMENT)")
     


# Reduce noise from HTTP libraries
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("azure.core.pipeline.policies.http_logging_policy").setLevel(logging.WARNING)

logger = logging.getLogger(__name__)


app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class Message(BaseModel):
    role: str
    content: str

class AskRequest(BaseModel):
    query: str
    prev_msgs: Optional[List[Message]] = None
    token: Optional[str] = None
    session_id: Optional[str] = None

print("Main module loaded successfully.")


# Helper function to encode base64
def encode_base64(text: str) -> str:
    if text is None:
        return None
    text_bytes = text.encode('utf-8')
    return base64.b64encode(text_bytes).decode('utf-8')



# Detect LibreOffice installation
def find_libreoffice():
    """
    Find LibreOffice executable on the system
    Returns path to soffice executable or None
    """
    possible_paths = [
        # Windows paths
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
        # Linux paths
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        # Mac paths
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            logger.info(f"Found LibreOffice at: {path}")
            return path
    
    # Try system PATH
    try:
        result = subprocess.run(['soffice', '--version'], capture_output=True, text=True)
        if result.returncode == 0:
            logger.info("Found LibreOffice in system PATH")
            return 'soffice'
    except:
        pass
    
    logger.warning("LibreOffice not found. PDF export with emojis may not work properly.")
    logger.warning("Install LibreOffice from: https://www.libreoffice.org/download/download/")
    return None


# Detect LibreOffice on startup
LIBREOFFICE_PATH = find_libreoffice()


def convert_docx_to_pdf_libreoffice(docx_path: str, output_dir: str) -> str:
    """
    Convert DOCX to PDF using LibreOffice
    This preserves emojis and formatting perfectly
    
    Args:
        docx_path: Path to input DOCX file
        output_dir: Directory for output PDF
    
    Returns:
        Path to generated PDF file
    """
    if not LIBREOFFICE_PATH:
        raise Exception("LibreOffice not found. Cannot convert DOCX to PDF. Please install LibreOffice.")
    
    try:
        cmd = [
            LIBREOFFICE_PATH,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ]
        
        logger.info(f"Running LibreOffice command: {' '.join(cmd)}")
        
        result = subprocess.run(
            cmd, 
            capture_output=True, 
            text=True, 
            timeout=30
        )
        
        if result.returncode != 0:
            logger.error(f"LibreOffice stderr: {result.stderr}")
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")
        
        # The PDF will be in output_dir with same name as DOCX
        docx_filename = os.path.basename(docx_path)
        pdf_filename = docx_filename.replace('.docx', '.pdf')
        pdf_path = os.path.join(output_dir, pdf_filename)
        
        if not os.path.exists(pdf_path):
            raise Exception(f"PDF file not created at expected path: {pdf_path}")
        
        logger.info(f"PDF created successfully: {pdf_path}")
        return pdf_path
        
    except subprocess.TimeoutExpired:
        raise Exception("LibreOffice conversion timed out (30s)")
    except Exception as e:
        logger.error(f"LibreOffice conversion error: {e}", exc_info=True)
        raise


# Cleanup function for expired items
def cleanup_expired_items():
    """Remove items older than 24 hours"""
    now = datetime.now(timezone.utc)
    expired_time = now - timedelta(hours=24)
    
    # Clean messages
    expired_msgs = [
        mid for mid, data in MESSAGES.items()
        if data.get("created_at", now) < expired_time
    ]
    for mid in expired_msgs:
        del MESSAGES[mid]
    
    # Clean tables
    expired_tables = [
        tid for tid, data in TABLES.items()
        if data.get("created_at", now) < expired_time
    ]
    for tid in expired_tables:
        del TABLES[tid]
    
    if expired_msgs or expired_tables:
        logger.info(f"Cleaned {len(expired_msgs)} expired messages and {len(expired_tables)} expired tables")


@app.post("/ask")
async def ask(
    request: Request,
    body: AskRequest = Body(...),
    encoded_string: str = Header(...)
):
    print(f"Received request body: {body}")

    query = body.query
    prev_msgs = body.prev_msgs or []

    # Initialize variables early
    database_name = None
    decrypted_fields = {}
    auth_token = None
    

    # Process encoded_string and generate auth_token first
    if encoded_string:
        try:
            decrypted = decode(encoded_string)
            print(f"Decrypted token: {decrypted}")
            database_name = decrypted.get("Database_Name")
            decrypted_fields = {
                "LoginMasterID": decrypted.get("LoginMasterID"),
                "Database_Name": decrypted.get("Database_Name"),
                "OrgID": decrypted.get("OrgID")
            }
            
            # Generate auth_token to call APIs
            if decrypted_fields:
                now_utc = datetime.now(timezone.utc)
                date_plus_one = (now_utc + timedelta(days=1)).isoformat()
                date_now = now_utc.isoformat()
                token_str = f"{date_plus_one}&{decrypted_fields.get('LoginMasterID')}&{decrypted_fields.get('Database_Name')}&{date_now}&{decrypted_fields.get('OrgID')}"
                auth_token = encode_base64(token_str)
                print(f"Base64 encoded auth_token to call api: {auth_token}")
                
        except Exception as e:
            logger.error(f"Failed to decode token: {e}")
    ###############################################################    
    #  DETECT DOWNLOAD REQUEST EARLY (before supervisor) to show download response link in UI if user explicitly asked for download in their query
    download_detection = detect_download_request(query)
    wants_download = download_detection.get("wants_download", False)
    format_preference = download_detection.get("format_preference", "any")
    download_message = download_detection.get("friendly_message", "")
    
    logger.info(f"Download detection: wants_download={wants_download}, format={format_preference}")
    
    # If user wants to download, find the previous assistant message and skip supervisor
    if wants_download and prev_msgs:
        previous_response = None
        original_question = None
        
        # Find the last assistant message and its corresponding user question
        for i in range(len(prev_msgs) - 1, -1, -1):
            msg = prev_msgs[i]
            msg_dict = msg.dict() if hasattr(msg, 'dict') else dict(msg) if not isinstance(msg, dict) else msg
            role = msg_dict.get('role', '')
            content = msg_dict.get('content', '')
            
            # Skip placeholder values from frontend
            if role in ['string', 'null', 'None', ''] or content in ['string', 'null', 'None', '']:
                continue
            
            # Found the assistant response
            if role == 'assistant' and not previous_response:
                previous_response = content
                # Now look backward to find the user question that came before it
                for j in range(i - 1, -1, -1):
                    prev_msg = prev_msgs[j]
                    prev_msg_dict = prev_msg.dict() if hasattr(prev_msg, 'dict') else dict(prev_msg) if not isinstance(prev_msg, dict) else prev_msg
                    prev_role = prev_msg_dict.get('role', '')
                    prev_content = prev_msg_dict.get('content', '')
                    
                    if prev_role in ['string', 'null', 'None', ''] or prev_content in ['string', 'null', 'None', '']:
                        continue
                    
                    if prev_role == 'user':
                        original_question = prev_content
                        break
                break
        
        if previous_response:
            logger.info(f"User wants to download previous response ({len(previous_response)} chars) - skipping supervisor")
            
            # Use previous response directly, skip supervisor call
            response_text = previous_response
            result = {"answer": response_text}
            
            # Clean expired items
            cleanup_expired_items()
            
            # Look up stored rows from TABLES dictionary by session ID
            rows = None
            skey = body.session_id or body.token
            if not skey or skey in ("null", "None", ""):
                skey = "sess_" + uuid4().hex[:12]
            
            # Search TABLES for most recent entry in this session
            for tid, tdata in sorted(TABLES.items(), key=lambda x: x[1].get("created_at", datetime.min), reverse=True):
                if tdata.get("session") == skey:
                    rows = tdata.get("rows")
                    if rows:
                        logger.info(f"Found stored rows from table {tid}: {len(rows)} rows")
                        break
            
            if not rows:
                logger.warning(f"No stored table data found for session {skey}")
            
            created_at = datetime.now(timezone.utc)
            expires_at = created_at + timedelta(hours=24)
            
            # Store the previous response for download with ORIGINAL question
            message_id = "msg_" + uuid4().hex[:10]
            MESSAGES[message_id] = {
                "question": original_question or query,  # Use original question, fallback to current query
                "answer": response_text,
                "session": skey,
                "created_at": created_at
            }
            
            # table_id = None
            # if has_markdown_table(response_text):
            #     # Try to extract table data from markdown if possible
            #     # For now, just mark that it has a table
            #     pass
            
            table_id = None
            # Create table_id if there are SQL rows (regardless of how they're displayed in response)
            if rows:
                table_id = "tbl_" + uuid4().hex[:10]
                TABLES[table_id] = {
                    "rows": rows,
                    "session": skey,
                    "created_at": created_at
                }
            
            # Get base URL
            base_url = str(request.url).replace("/ask", "")
            if base_url.startswith("http://"):
                base_url = base_url.replace("http://", "https://", 1)
            
            # CONDITIONALLY GENERATE DOWNLOAD URLS BASED ON FORMAT PREFERENCE AND DATA AVAILABILITY
            download_urls = {}
            export_info = None
            
            # Excel - ONLY if user requested data download AND there's a table
            if table_id and format_preference in ["excel", "both", "any"]:
                download_urls["excel"] = f"{base_url}/export/table/{table_id}"
            
            # DOCX - only if user requested response download (but NOT if they specifically want Excel only)
            if message_id and format_preference in ["docx", "both", "any"]:
                download_urls["docx"] = f"{base_url}/export/message/{message_id}?format=docx"
            
            # If user wants Excel but no table exists, offer DOCX as fallback
            if format_preference == "excel" and not table_id:
                # Provide DOCX as fallback
                if message_id:
                    download_urls["docx"] = f"{base_url}/export/message/{message_id}?format=docx"
                    friendly_response = "I don't have tabular data from the previous response to export as Excel, but I've prepared a Word document with the response for you to download."
                    logger.warning("User requested Excel download but no table data available - offering DOCX fallback")
                else:
                    friendly_response = "I don't have any data from the previous response to export. Please ask a question first, then request the download."
            else:
                # Use the friendly message from download detector as the answer
                friendly_response = download_message
            
            print(f"response from download detector: {friendly_response}")
            
            # Only create export_info if we have download URLs
            if download_urls:
                export_info = {
                    "message_id": message_id,
                    "table_id": table_id,
                    "download_urls": download_urls,
                    "storage_type": "https",
                    "expires_in": "24 hours",
                    "expires_at": expires_at.isoformat()
                }
            
            timestamp_bot = datetime.now(timezone.utc).isoformat()
            context_list = [
                {"role": "user", "content": query, "timestamp": timestamp_bot},
                {"role": "assistant", "content": friendly_response, "timestamp": timestamp_bot}
            ]
            user_details = {
                "session_id": getattr(body, "session_id", None) if hasattr(body, "session_id") else None,
                "token": getattr(body, "token", None) if hasattr(body, "token") else None
            }
            
            return {
                "answer": friendly_response,
                "timestamp": timestamp_bot,
                "context": context_list,
                "user_details": user_details,
                "decrypted_fields": decrypted_fields,
                "export": export_info
            }
        else:
            logger.warning("User wants to download but no previous assistant message found")
            # Return error message instead of calling supervisor
            timestamp_bot = datetime.now(timezone.utc).isoformat()
            return {
                "answer": "There's no previous response to download. Please ask a question first.",
                "timestamp": timestamp_bot,
                "context": [
                    {"role": "user", "content": query, "timestamp": timestamp_bot},
                    {"role": "assistant", "content": "There's no previous response to download. Please ask a question first.", "timestamp": timestamp_bot}
                ],
                "user_details": {
                    "session_id": getattr(body, "session_id", None) if hasattr(body, "session_id") else None,
                    "token": getattr(body, "token", None) if hasattr(body, "token") else None
                },
                "decrypted_fields": decrypted_fields
            }

   
    #########################################################

    # Now rewrite the question using contextllm with auth_token available
    try:
        prev_msgs_dict = [m.dict() if hasattr(m, 'dict') else dict(m) if not isinstance(m, dict) else m for m in prev_msgs]
        full_question = rewrite_question(prev_msgs_dict, query, auth_token)
        print(f"[contextllm] Rewritten full question: {full_question}")
        logger.info(f"[contextllm] Rewritten full question: {full_question}")
    except Exception as e:
        print(f"[contextllm] Failed to rewrite question, using fallback. Error: {e}")
        logger.error(f"[contextllm] Failed to rewrite question, using fallback. Error: {e}")
        last_msgs = prev_msgs[-5:]
        context = "\n".join([f"Previous message {i+1} ({msg['role'] if isinstance(msg, dict) else msg.role}): {msg['content'] if isinstance(msg, dict) else msg.content}" for i, msg in enumerate(last_msgs)])
        full_question = f"{context}\nCurrent question: {query}" if context else query
        print(f"Full question: {full_question}")
        logger.info(f"Full question: {full_question}")

    result = await supervisor(full_question, database_name, auth_token)
    # print(f"Printing final result: {result}")

    # Extract the main response text
    response_text = None
    if isinstance(result, dict):
        if "answer" in result:
            response_text = result["answer"]
            while isinstance(response_text, dict) and "answer" in response_text:
                response_text = response_text["answer"]
        elif "error" in result:
            response_text = "Server is down, please try again in some time."
        else:
            response_text = str(result)
    else:
        response_text = str(result)

    if not isinstance(response_text, str):
        response_text = json.dumps(response_text, ensure_ascii=False)
    
    # Log the full text length for debugging
    logger.info(f"Storing message with {len(response_text)} characters")
        
    # Clean expired items before creating new ones
    cleanup_expired_items()
    
    rows = result.get("rows") if isinstance(result, dict) else None
    if not isinstance(rows, list):
        rows = None

    sql_query = result.get("sql_query") if isinstance(result, dict) else None

    skey = body.session_id or body.token
    if not skey or skey in ("null", "None", ""):
        skey = "sess_" + uuid4().hex[:12]
    
    # Create timestamp for expiration
    created_at = datetime.now(timezone.utc)
    expires_at = created_at + timedelta(hours=24)
    
    # Create new IDs + store with timestamp (stores original text with emojis)
    message_id = "msg_" + uuid4().hex[:10]
    MESSAGES[message_id] = {
        "question": query,   # original user question
        "answer": response_text,  # Original response with emojis preserved
        "session": skey,
        "created_at": created_at
    }

    
    # table_id = None
    # # Only create table_id if there are rows AND there's a markdown table in the response
    # if rows and has_markdown_table(response_text):
    #     table_id = "tbl_" + uuid4().hex[:10]
    #     TABLES[table_id] = {
    #         "rows": rows,
    #         "session": skey,
    #         "created_at": created_at
    #     }
    
    table_id = None
    # Create table_id if there are SQL rows (regardless of how they're displayed in response)
    if rows:
        table_id = "tbl_" + uuid4().hex[:10]
        TABLES[table_id] = {
            "rows": rows,
            "session": skey,
            "created_at": created_at
        }
        
    # Get base URL from request
    # base_url = str(request.url).replace("/ask", "")
        # Get base URL from request and force HTTPS
    base_url = str(request.url).replace("/ask", "")
    if base_url.startswith("http://"):
        base_url = base_url.replace("http://", "https://", 1)
    
    
    # # Generate download URLs
    # download_urls = {}
    # if message_id:
    #     download_urls["docx"] = f"{base_url}/export/message/{message_id}?format=docx"
    #     download_urls["pdf"] = f"{base_url}/export/message/{message_id}?format=pdf"
    
    # if table_id:
    #     download_urls["excel"] = f"{base_url}/export/table/{table_id}"

    # timestamp_bot = datetime.now(timezone.utc).isoformat()
    # context_list = [
    #     {"role": "user", "content": query, "timestamp": timestamp_bot},
    #     {"role": "assistant", "content": response_text, "timestamp": timestamp_bot}
    # ]
    # user_details = {
    #     "session_id": getattr(body, "session_id", None) if hasattr(body, "session_id") else None,
    #     "token": getattr(body, "token", None) if hasattr(body, "token") else None
    # }

    # return {
    #     "answer": response_text,  # Original response with emojis
    #     "timestamp": timestamp_bot,
    #     "context": context_list,
    #     "user_details": user_details,
    #     "export": {
    #         "message_id": message_id,
    #         "table_id": table_id,
    #         "download_urls": download_urls,
    #         "storage_type": "https",
    #         "expires_in": "24 hours",
    #         "expires_at": expires_at.isoformat()
    #     },  
    #     "decrypted_fields": decrypted_fields
    # }
    
    
    # CONDITIONALLY GENERATE DOWNLOAD URLS BASED ON USER REQUEST - ONLY IF USER EXPLICITLY ASKED FOR DOWNLOAD IN THEIR QUERY (wants_download=True from download_detector)
    download_urls = {}
    export_info = None
    
    # DOCX - only if user requested download
    # if wants_download and message_id:
    #     download_urls["docx"] = f"{base_url}/export/message/{message_id}?format=docx"
    # DOCX - only if user requested response download
    if wants_download and message_id:
        if format_preference in ["docx", "both", "any"]:
            download_urls["docx"] = f"{base_url}/export/message/{message_id}?format=docx"
    
    
    # # Excel - ALWAYS available when there's a table (unchanged behavior)
    # if table_id:
    #     download_urls["excel"] = f"{base_url}/export/table/{table_id}"
    
     # Excel - ONLY if user requested data download AND there's a table
    if wants_download and table_id:
        if format_preference in ["excel", "both", "any"]:
            download_urls["excel"] = f"{base_url}/export/table/{table_id}"
    
    # Only create export_info if we have download URLs
    if download_urls:
        export_info = {
            "message_id": message_id,
            "table_id": table_id,
            "download_urls": download_urls,
            "storage_type": "https",
            "expires_in": "24 hours",
            "expires_at": expires_at.isoformat()
        }

    timestamp_bot = datetime.now(timezone.utc).isoformat()
    context_list = [
        {"role": "user", "content": query, "timestamp": timestamp_bot},
        {"role": "assistant", "content": response_text, "timestamp": timestamp_bot}
    ]
    user_details = {
        "session_id": getattr(body, "session_id", None) if hasattr(body, "session_id") else None,
        "token": getattr(body, "token", None) if hasattr(body, "token") else None
    }

    response_data = {
        "answer": response_text,  # Clean response WITHOUT download links
        "timestamp": timestamp_bot,
        "context": context_list,
        "user_details": user_details,
        "decrypted_fields": decrypted_fields,
        **({"sql_query": sql_query} if sql_query else {})
    }
    
    # Only include export info if we have download URLs
    if export_info:
        response_data["export"] = export_info

    return response_data

    


@app.get("/export/message/{message_id}")
async def export_message(message_id: str, format: str = Query("docx")):
    """
    Export stored message to DOCX or PDF format
    PDF is generated by converting DOCX (preserves emojis perfectly)
    """
    item = MESSAGES.get(message_id)
    if not item:
        raise HTTPException(status_code=404, detail="Message not found or has expired")
    
    created_at = item.get("created_at", datetime.now(timezone.utc))
    expires_at = created_at + timedelta(hours=24)
    if datetime.now(timezone.utc) > expires_at:
        del MESSAGES[message_id]
        raise HTTPException(status_code=410, detail="Download link has expired (24 hours)")

    question = item.get("question", "")
    text = item["answer"] or ""
    fmt = (format or "docx").lower()

    logger.info(f"Exporting message {message_id} with {len(text)} characters in format {fmt}")

    # Generate DOCX first (used for both DOCX and PDF exports)
    doc = Document()
    
    # Adding the question as a heading or section
    if question:
        doc.add_heading('User Question:', level=1)
        doc.add_paragraph(question)
        doc.add_heading('AI Response:', level=1)
    
    try:
        html = markdown.markdown(text, extensions=['tables', 'nl2br', 'fenced_code'])
        soup = BeautifulSoup(html, 'html.parser')
        
        has_structure = bool(soup.find_all(['h1', 'h2', 'h3', 'ul', 'ol', 'table']))
        
        if has_structure:
            for element in soup.find_all(recursive=False):
                if element.name in ['h1', 'h2', 'h3']:
                    doc.add_heading(element.get_text(), level=int(element.name[1]))
                elif element.name == 'p':
                    doc.add_paragraph(element.get_text())
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.get_text(), style='List Bullet')
                elif element.name == 'ol':
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.get_text(), style='List Number')
                elif element.name == 'table':
                    rows = []
                    for tr in element.find_all('tr'):
                        row = [td.get_text().strip() for td in tr.find_all(['td', 'th'])]
                        if row:
                            rows.append(row)
                    
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Light Grid Accent 1'
                        for i, row_data in enumerate(rows):
                            for j, cell_data in enumerate(row_data):
                                table.rows[i].cells[j].text = cell_data
        else:
            # Plain text - preserve emojis
            for line in text.split('\n'):
                doc.add_paragraph(line)
    except Exception as e:
        logger.error(f"Error in DOCX markdown conversion: {e}")
        # Fallback: plain text
        for line in text.split('\n'):
            doc.add_paragraph(line)

    # If DOCX format requested, return directly
    if fmt == "docx":
        docx_buf = io.BytesIO()
        doc.save(docx_buf)
        docx_buf.seek(0)

        return Response(
            content=docx_buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="response.docx"'}
        )

    # If PDF format requested, convert DOCX to PDF using LibreOffice
    if fmt == "pdf":
        if not LIBREOFFICE_PATH:
            raise HTTPException(
                status_code=503, 
                detail="PDF export unavailable. LibreOffice not installed. Please install from https://www.libreoffice.org/"
            )
        
        temp_dir = None
        try:
            # Create temporary directory
            temp_dir = tempfile.mkdtemp()
            logger.info(f"Created temp directory: {temp_dir}")
            
            # Save DOCX to temp file
            temp_docx_path = os.path.join(temp_dir, f"{message_id}.docx")
            doc.save(temp_docx_path)
            logger.info(f"Saved DOCX to: {temp_docx_path}")
            
            # Convert DOCX to PDF using LibreOffice
            logger.info("Converting DOCX to PDF using LibreOffice...")
            pdf_path = convert_docx_to_pdf_libreoffice(temp_docx_path, temp_dir)
            
            # Read the PDF content
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            logger.info(f"PDF generated successfully with {len(pdf_content)} bytes")
            
            return Response(
                content=pdf_content,
                media_type="application/pdf",
                headers={"Content-Disposition": 'attachment; filename="response.pdf"'}
            )
            
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {e}", exc_info=True)
            raise HTTPException(
                status_code=500, 
                detail=f"Failed to generate PDF: {str(e)}"
            )
        
        finally:
            # Clean up temporary directory
            if temp_dir and os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                    logger.info(f"Cleaned up temp directory: {temp_dir}")
                except Exception as e:
                    logger.warning(f"Failed to delete temp directory: {e}")
    
    raise HTTPException(status_code=400, detail="format must be docx or pdf")


@app.get("/export/table/{table_id}")
async def export_table(table_id: str):
    """Export stored table data to Excel format"""
    item = TABLES.get(table_id)
    if not item:
        raise HTTPException(status_code=404, detail="Table not found or has expired")
    
    created_at = item.get("created_at", datetime.now(timezone.utc))
    expires_at = created_at + timedelta(hours=24)
    if datetime.now(timezone.utc) > expires_at:
        del TABLES[table_id]
        raise HTTPException(status_code=410, detail="Download link has expired (24 hours)")

    rows = item["rows"]

    df = pd.DataFrame(rows)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Data")

    buf.seek(0)
    filename = "table_export.xlsx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )
